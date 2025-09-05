import os
import PyPDF2
import docx
import io
from typing import Optional

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from PDF or DOCX files
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    except Exception as e:
        print(f"Text extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from file: {str(e)}")

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF using PyPDF2"""

    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise Exception("PDF is password protected")

            # Extract text from all pages
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if not text.strip():
            raise Exception("No text could be extracted from PDF")

        return clean_extracted_text(text)

    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX using python-docx"""

    try:
        doc = docx.Document(docx_path)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        if not text.strip():
            raise Exception("No text could be extracted from DOCX")

        return clean_extracted_text(text)

    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""

    if not text:
        return ""

    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        cleaned_line = ' '.join(line.split())  # Remove extra spaces
        if cleaned_line:  # Only add non-empty lines
            cleaned_lines.append(cleaned_line)

    # Join lines with single newline
    cleaned_text = '\n'.join(cleaned_lines)

    # Remove multiple consecutive newlines
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')

    return cleaned_text.strip()

def extract_basic_info(text: str) -> dict:
    """Extract basic information using simple patterns"""

    import re

    info = {
        'emails': [],
        'phones': [],
        'urls': []
    }

    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    info['emails'] = re.findall(email_pattern, text)

    # Phone pattern (various formats)
    phone_patterns = [
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # 123-456-7890 or 123.456.7890
        r'\(\d{3}\)\s?\d{3}[-.]?\d{4}',    # (123) 456-7890
        r'\b\d{10}\b'                         # 1234567890
    ]

    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        info['phones'].extend(phones)

    # URL pattern
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    info['urls'] = re.findall(url_pattern, text)

    # LinkedIn profile pattern
    linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9-]+'
    linkedin_profiles = re.findall(linkedin_pattern, text.lower())
    info['linkedin'] = linkedin_profiles

    # GitHub profile pattern
    github_pattern = r'github\.com/[A-Za-z0-9-]+'
    github_profiles = re.findall(github_pattern, text.lower())
    info['github'] = github_profiles

    return info

# Alternative PDF extraction method using pdfplumber (if PyPDF2 fails)
def extract_text_from_pdf_alternative(pdf_path: str) -> str:
    """Alternative PDF extraction using pdfplumber"""

    try:
        import pdfplumber

        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return clean_extracted_text(text)

    except ImportError:
        raise Exception("pdfplumber not installed. Using PyPDF2 fallback.")
    except Exception as e:
        raise Exception(f"Alternative PDF extraction failed: {str(e)}")
